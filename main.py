import os
import re

from bs4 import BeautifulSoup, Tag, ResultSet
from docx import Document
from docx.shared import RGBColor
from docx.table import Table


def extract_values(file_path: str) -> dict[str, int]:
    with open(f"data/{file_path}", "r", encoding="utf-8") as file:
        soup: BeautifulSoup = BeautifulSoup(file, "html.parser")
        text: str = soup.get_text()

        match: re.Match[str] = re.search(r"F\(X\)\s*=\s*([\d\+\-\*x\s]+)", text)
        if match:
            function_str: str = match.group(1)
            coefficients: list[tuple[str]] = re.findall(r"(\d+)\s*x(\d+)", function_str)
            return {f"x{var}": int(coeff) for coeff, var in coefficients}


def extract_tables(file_path: str) -> list[Tag]:
    with open(f"data/{file_path}", "r", encoding="utf-8") as file:
        soup: BeautifulSoup = BeautifulSoup(file, "html.parser")
        tables: ResultSet = soup.find_all(
            "table", class_="table table-bordered table-center"
        )
        filtered_tables: list[Tag] = []

        if tables:
            for table in tables:
                if table.find("td", bgcolor="FFA0A0"):
                    filtered_tables.append(table)

            if tables[-1] not in filtered_tables:
                filtered_tables.append(tables[-1])

        return filtered_tables


def transform_tables(table: Tag, target_values: dict[str, int]) -> Tag:
    # Добавляем строку сверху
    num_columns: int = len(table.find_all("tr")[0].find_all(["td", "th"]))
    header_row: Tag = Tag(name="tr")

    c_cell: Tag = Tag(name="td")
    c_cell.string = "C"
    header_row.append(c_cell)

    dash_cell: Tag = Tag(name="td")
    dash_cell.string = "-"
    header_row.append(dash_cell)

    for i in range(num_columns - 2):
        cell: Tag = Tag(name="td")
        cell.string = str(target_values.get(f"x{i + 1}", 0))
        header_row.append(cell)

    table.insert(0, header_row)

    # Добавляем столбец слева
    rows_len: int = len(table.find_all("tr"))
    for i, row in enumerate(table.find_all("tr")):
        first_cell: Tag = Tag(name="td")
        if i == 0 or i == 1 or i == rows_len - 1:
            first_cell.string = ""
        else:
            variable_name: str = row.find_all("td")[0].get_text().strip()
            first_cell.string = str(target_values.get(variable_name, 0))
        row.insert(0, first_cell)

    return table


def extract_result(file_path: str) -> str:
    with open(f"data/{file_path}", "r", encoding="utf-8") as file:
        soup: BeautifulSoup = BeautifulSoup(file, "html.parser")
        text: str = soup.get_text()

        optimal_plan_pattern: re.Pattern[str] = re.compile(
            r"F\(X\)\s*=\s*[0-9*\/\+\-\s]+=\s*[0-9*\/]+", re.DOTALL
        )

        match: re.Match[str] = optimal_plan_pattern.search(text)
        if match:
            return f"Результат: {match.group(0)}"
        else:
            return (
                f"Результат: Последняя строка содержит отрицательные элементы. "
                "Пространство допустимых решений неограниченно. Решения не существует"
            )


def save_table_to_docx(
    tables: list[Tag],
    output_file: str,
    result_text: str,
) -> None:
    doc: Document = Document()

    for table in tables:
        doc_table: Table = doc.add_table(
            rows=0, cols=len(table.find_all("tr")[0].find_all(["td", "th"]))
        )

        for row in table.find_all("tr"):
            doc_row = doc_table.add_row().cells
            for i, cell in enumerate(row.find_all("td")):
                doc_row[i].text = cell.get_text().strip()
                if cell.get("bgcolor"):
                    run = doc_row[i].paragraphs[0].runs[0]
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True

    doc.add_paragraph(result_text)
    doc.save(output_file)


def parse():
    html_files: list[str] = ["data1.html", "data2.html", "data3.html"]
    output_dir: str = "extracted"
    os.makedirs(output_dir, exist_ok=True)

    for html_file in html_files:
        target_values: dict[str, int] = extract_values(html_file)
        tables: list[Tag] = extract_tables(html_file)
        result: str = extract_result(html_file)

        transformed_tables: list[Tag] = [
            transform_tables(table, target_values) for table in tables
        ]
        output_docx_path: str = os.path.join(
            output_dir, f'{os.path.basename(html_file).split(".")[0]}.docx'
        )
        save_table_to_docx(transformed_tables, output_docx_path, result)


if __name__ == "__main__":
    parse()
